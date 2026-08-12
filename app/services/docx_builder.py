from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "063B67"
BLUE = "92B2D8"
LIGHT_BLUE = "E9F1F8"
PALE_BLUE = "F5F8FC"
YELLOW = "F7C600"
GRAY = "D9D9D9"
DARK_GRAY = "333333"
WHITE = "FFFFFF"
RED = "9C2F2F"


@dataclass(slots=True)
class GeneratedArtifact:
    filename: str
    mimetype: str
    content: io.BytesIO


def _slug(value: str, fallback: str = "documento") -> str:
    normalized = re.sub(r"[^A-Za-z0-9À-ÿ_-]+", "_", value.strip())
    normalized = re.sub(r"_+", "_", normalized).strip("_")
    return (normalized or fallback)[:80]


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = "555555", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_run_font(run, size: float = 10, bold: bool | None = None, color: str = DARK_GRAY, italic: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _format_paragraph(
    paragraph,
    *,
    size: float = 10,
    bold: bool | None = None,
    color: str = DARK_GRAY,
    alignment=None,
    space_before: float = 0,
    space_after: float = 3,
    line_spacing: float = 1.05,
) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        _set_run_font(run, size=size, bold=bold, color=color)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _write_cell(
    cell,
    text: str,
    *,
    size: float = 9,
    bold: bool = False,
    color: str = DARK_GRAY,
    fill: str | None = None,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    _set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text or ""))
    _set_run_font(run, size=size, bold=bold, color=color)


def _write_labeled_cell(cell, label: str, value: str, *, size: float = 9, fill: str | None = None) -> None:
    _set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(label)
    _set_run_font(run, size=size, bold=True, color=NAVY)
    run = paragraph.add_run(value or "—")
    _set_run_font(run, size=size, color=DARK_GRAY)


def _add_bullets(container, items: Iterable[str], *, size: float = 9.5, compact: bool = True) -> None:
    for item in items or []:
        if not str(item).strip():
            continue
        paragraph = container.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Cm(0.45)
        paragraph.paragraph_format.first_line_indent = Cm(-0.2)
        paragraph.paragraph_format.space_after = Pt(1.5 if compact else 4)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(str(item).strip())
        _set_run_font(run, size=size)


def _add_numbered(container, items: Iterable[str], *, size: float = 10) -> None:
    for item in items or []:
        if not str(item).strip():
            continue
        paragraph = container.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(str(item).strip())
        _set_run_font(run, size=size)


def _usable_width(section) -> int:
    return section.page_width - section.left_margin - section.right_margin


def _configure_document(doc: Document, *, landscape: bool = False, body_size: float = 10) -> None:
    section = doc.sections[0]
    section.page_width = Cm(29.7 if landscape else 21.0)
    section.page_height = Cm(21.0 if landscape else 29.7)
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 13, NAVY), ("Heading 3", 11, DARK_GRAY)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    doc.core_properties.title = "Documento institucional gerado para revisão docente"
    doc.core_properties.author = "Gerador Institucional de Documentos"
    doc.core_properties.subject = "Planejamento e documentação acadêmica"
    doc.core_properties.comments = "Conteúdo gerado por IA; requer validação e aprovação institucional."


def _add_header_footer(doc: Document, config: dict[str, Any]) -> None:
    logo_path = Path(config.get("LOGO_PATH", "")) if config.get("LOGO_PATH") else None
    for section in doc.sections:
        header = section.header
        table = header.add_table(rows=1, cols=2, width=_usable_width(section))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_cell_width(table.cell(0, 0), 8.5)
        _set_cell_width(table.cell(0, 1), 9.0 if section.orientation == WD_ORIENT.PORTRAIT else 17.0)
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        _set_cell_margins(left, 0, 0, 0, 0)
        _set_cell_margins(right, 0, 0, 0, 0)
        if logo_path and logo_path.exists():
            paragraph = left.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.add_run().add_picture(str(logo_path), width=Cm(3.3))
        else:
            _write_cell(left, config.get("INSTITUTION_NAME", "Instituição de Ensino Superior"), size=10, bold=True, color=NAVY)

        paragraph = right.paragraphs[0]
        _clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for index, value in enumerate(
            filter(
                None,
                [
                    config.get("INSTITUTION_UNIT", ""),
                    config.get("INSTITUTION_WEBSITE", ""),
                    config.get("INSTITUTION_ADDRESS", ""),
                ],
            )
        ):
            if index:
                paragraph.add_run("\n")
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=7.5 if index else 8.5, bold=index == 0, color=NAVY)

        # Linha fina sob o cabeçalho.
        p = header.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), BLUE)
        borders.append(bottom)
        p_pr.append(borders)

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = config.get("INSTITUTION_FOOTER") or "Documento sujeito à revisão e aprovação docente/institucional."
        run = paragraph.add_run(footer_text)
        _set_run_font(run, size=7, color=NAVY)


def _add_title(doc: Document, title: str, subtitle: str = "") -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title.upper())
    _set_run_font(run, size=15, bold=True, color=NAVY)
    if subtitle:
        run = paragraph.add_run(f"\n{subtitle}")
        _set_run_font(run, size=9, color=DARK_GRAY)


def _add_section_box(doc: Document, title: str, body: str | list[str], *, bullets: bool = False) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table, color="7F8C97", size=5)
    _write_cell(table.cell(0, 0), title.upper(), size=9, bold=True, color=NAVY, fill=LIGHT_BLUE)
    cell = table.cell(1, 0)
    _set_cell_margins(cell, 110, 130, 110, 130)
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    if bullets:
        _add_bullets(cell, body if isinstance(body, list) else [body], size=9.5)
    else:
        items = body if isinstance(body, list) else [body]
        for index, item in enumerate(items):
            p = paragraph if index == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(str(item))
            _set_run_font(run, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_review_notes(doc: Document, notes: list[str], title: str = "PONTOS PARA REVISÃO DOCENTE") -> None:
    if not notes:
        return
    table = doc.add_table(rows=1, cols=1)
    _set_table_borders(table, color="D8A000", size=6)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "FFF7D6")
    _set_cell_margins(cell, 130, 150, 130, 150)
    p = cell.paragraphs[0]
    _clear_paragraph(p)
    run = p.add_run(title)
    _set_run_font(run, size=9, bold=True, color="7A5900")
    _add_bullets(cell, notes, size=8.8)


def _save(doc: Document) -> io.BytesIO:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_ementa(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    doc = Document()
    _configure_document(doc, body_size=10)
    _add_header_footer(doc, config)
    _add_title(doc, "Ementário")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)
    _write_labeled_cell(table.cell(0, 0), "COMPONENTE CURRICULAR: ", form.get("componente", ""), size=9, fill=GRAY)
    _write_labeled_cell(table.cell(0, 1), "CARGA HORÁRIA: ", form.get("carga_horaria", ""), size=9, fill=GRAY)

    ementa_cell = table.cell(1, 0).merge(table.cell(1, 1))
    _set_cell_margins(ementa_cell, 120, 140, 120, 140)
    p = ementa_cell.paragraphs[0]
    _clear_paragraph(p)
    run = p.add_run("EMENTA: ")
    _set_run_font(run, size=9.5, bold=True, color=NAVY)
    run = p.add_run(data.get("ementa", ""))
    _set_run_font(run, size=9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.08

    _write_cell(table.cell(2, 0), "BIBLIOGRAFIA BÁSICA", size=9, bold=True, color=NAVY, fill=GRAY)
    _write_cell(table.cell(2, 1), "BIBLIOGRAFIA COMPLEMENTAR", size=9, bold=True, color=NAVY, fill=GRAY)
    for column, key in ((0, "bibliografia_basica"), (1, "bibliografia_complementar")):
        cell = table.cell(3, column)
        _set_cell_margins(cell, 120, 140, 120, 140)
        _clear_paragraph(cell.paragraphs[0])
        _add_bullets(cell, data.get(key, []), size=9, compact=False)

    _add_review_notes(doc, data.get("notas_para_revisao", []))
    filename = f"Ementa_{_slug(form.get('componente', 'componente'))}.docx"
    return GeneratedArtifact(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _save(doc))


def build_plano_ensino(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    doc = Document()
    _configure_document(doc, body_size=9.5)
    _add_header_footer(doc, config)
    _add_title(doc, "Plano de Ensino", form.get("semestre_letivo", ""))

    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _set_table_borders(table)
    title_cell = table.cell(0, 0).merge(table.cell(0, 3))
    _write_cell(title_cell, "PLANO DE ENSINO", size=9.5, bold=True, color=NAVY, fill=LIGHT_BLUE)
    course_cell = table.cell(1, 0).merge(table.cell(1, 1))
    component_cell = table.cell(1, 2).merge(table.cell(1, 3))
    _write_labeled_cell(course_cell, "Curso: ", form.get("curso", ""))
    _write_labeled_cell(component_cell, "Componente: ", form.get("componente", ""))
    _write_labeled_cell(table.cell(2, 0), "Período: ", form.get("periodo", ""))
    _write_labeled_cell(table.cell(2, 1), "Turno: ", form.get("turno", ""))
    _write_labeled_cell(table.cell(2, 2), "Modalidade: ", form.get("modalidade", ""))
    _write_labeled_cell(table.cell(2, 3), "Semestre: ", form.get("semestre_letivo", ""))
    professor_cell = table.cell(3, 0).merge(table.cell(3, 1))
    _write_labeled_cell(professor_cell, "Professor responsável: ", form.get("professor_responsavel", ""))
    _write_labeled_cell(table.cell(3, 2), "C/H semanal: ", form.get("ch_semanal", ""))
    _write_labeled_cell(table.cell(3, 3), "C/H semestral: ", form.get("ch_semestral", ""))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    _add_section_box(doc, "Ementa", data.get("ementa", ""))

    objectives = data.get("objetivos", {})
    obj_table = doc.add_table(rows=4, cols=2)
    obj_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(obj_table)
    heading = obj_table.cell(0, 0).merge(obj_table.cell(0, 1))
    _write_cell(heading, "OBJETIVOS / COMPETÊNCIAS", size=9, bold=True, color=NAVY, fill=LIGHT_BLUE)
    for row_index, (label, key) in enumerate(
        (("Cognitivos", "cognitivos"), ("Habilidades", "habilidades"), ("Atitudes", "atitudes")), start=1
    ):
        _write_cell(obj_table.cell(row_index, 0), label, size=9, bold=True, color=NAVY, fill=PALE_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell = obj_table.cell(row_index, 1)
        _set_cell_margins(cell)
        _clear_paragraph(cell.paragraphs[0])
        _add_bullets(cell, objectives.get(key, []), size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    units = data.get("unidades", [])
    content_table = doc.add_table(rows=2 + len(units), cols=3)
    content_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(content_table)
    heading = content_table.cell(0, 0).merge(content_table.cell(0, 2))
    _write_cell(heading, "CONTEÚDO", size=9, bold=True, color=NAVY, fill=LIGHT_BLUE)
    for index, title in enumerate(("UNIDADE", "C/H", "TÓPICOS")):
        _write_cell(content_table.cell(1, index), title, size=8.5, bold=True, color=NAVY, fill=PALE_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_table_header(content_table.rows[1])
    for row_index, unit in enumerate(units, start=2):
        _write_cell(content_table.cell(row_index, 0), unit.get("unidade", ""), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(content_table.cell(row_index, 1), unit.get("carga_horaria", ""), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cell = content_table.cell(row_index, 2)
        _set_cell_margins(cell)
        _clear_paragraph(cell.paragraphs[0])
        _add_bullets(cell, unit.get("topicos", []), size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    _add_section_box(doc, "Estratégias de ensino", data.get("estrategia_ensino", []), bullets=True)
    _add_section_box(doc, "Recursos disponíveis", data.get("recursos_disponiveis", []), bullets=True)
    _add_section_box(doc, "Avaliação", data.get("avaliacao", []), bullets=True)

    bib_table = doc.add_table(rows=2, cols=2)
    bib_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(bib_table)
    for index, title in enumerate(("BIBLIOGRAFIA BÁSICA", "BIBLIOGRAFIA COMPLEMENTAR")):
        _write_cell(bib_table.cell(0, index), title, size=8.8, bold=True, color=NAVY, fill=LIGHT_BLUE)
    for column, key in ((0, "bibliografia_basica"), (1, "bibliografia_complementar")):
        cell = bib_table.cell(1, column)
        _set_cell_margins(cell)
        _clear_paragraph(cell.paragraphs[0])
        _add_bullets(cell, data.get(key, []), size=8.5, compact=False)

    _add_review_notes(doc, data.get("notas_para_revisao", []))
    filename = f"Plano_de_Ensino_{_slug(form.get('componente', 'componente'))}_{_slug(form.get('semestre_letivo', 'semestre'))}.docx"
    return GeneratedArtifact(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _save(doc))


def build_cronograma(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    doc = Document()
    _configure_document(doc, landscape=True, body_size=8.5)
    _add_header_footer(doc, config)
    _add_title(doc, "Cronograma de Aulas", form.get("semestre_letivo", ""))

    meta = doc.add_table(rows=3, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(meta)
    _write_labeled_cell(meta.cell(0, 0).merge(meta.cell(0, 1)), "Curso: ", form.get("curso", ""))
    _write_labeled_cell(meta.cell(0, 2).merge(meta.cell(0, 3)), "Componente: ", form.get("componente", ""))
    _write_labeled_cell(meta.cell(1, 0), "Período: ", form.get("periodo", ""))
    _write_labeled_cell(meta.cell(1, 1), "Turma: ", form.get("turma", ""))
    _write_labeled_cell(meta.cell(1, 2), "Semestre: ", form.get("semestre_letivo", ""))
    _write_labeled_cell(meta.cell(1, 3), "Coordenação: ", form.get("coordenador", ""))
    docentes = ", ".join(form.get("docentes", []))
    _write_labeled_cell(meta.cell(2, 0).merge(meta.cell(2, 1)), "Docente(s): ", docentes)
    _write_labeled_cell(meta.cell(2, 2), "CH teórica: ", form.get("ch_teorica", ""))
    _write_labeled_cell(meta.cell(2, 3), "CH prática: ", form.get("ch_pratica", ""))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    headers = ["Aula", "Data", "Dia", "Horário", "C/H", "Natureza", "Tópico", "Docente", "Observações"]
    rows = data.get("aulas", [])
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, color="6A7782", size=5)
    widths = [1.2, 2.0, 2.2, 2.5, 1.4, 2.5, 7.4, 3.7, 3.4]
    for index, (header, width) in enumerate(zip(headers, widths)):
        _set_cell_width(table.cell(0, index), width)
        _write_cell(table.cell(0, index), header, size=7.5, bold=True, color=WHITE, fill=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_table_header(table.rows[0])
    for row_index, item in enumerate(rows, start=1):
        values = [
            item.get("numero", ""), item.get("data", ""), item.get("dia_semana", ""), item.get("horario", ""),
            item.get("carga_horaria", ""), item.get("natureza", ""), item.get("topico", ""),
            item.get("docente", ""), item.get("observacoes", ""),
        ]
        for column, (value, width) in enumerate(zip(values, widths)):
            _set_cell_width(table.cell(row_index, column), width)
            align = WD_ALIGN_PARAGRAPH.LEFT if column in {6, 7, 8} else WD_ALIGN_PARAGRAPH.CENTER
            fill = PALE_BLUE if row_index % 2 == 0 else None
            _write_cell(table.cell(row_index, column), str(value), size=7.2, fill=fill, alignment=align)

    summary = [
        f"Carga horária teórica planejada: {data.get('carga_horaria_teorica_planejada', '')}",
        f"Carga horária prática planejada: {data.get('carga_horaria_pratica_planejada', '')}",
    ]
    _add_section_box(doc, "Conferência de carga horária", summary, bullets=True)
    _add_review_notes(doc, data.get("alertas_calendario", []) + data.get("notas_para_revisao", []))
    filename = f"Cronograma_{_slug(form.get('componente', 'componente'))}_{_slug(form.get('turma', 'turma'))}.docx"
    return GeneratedArtifact(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _save(doc))


def _build_single_lesson_plan(plan: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> io.BytesIO:
    doc = Document()
    _configure_document(doc, body_size=9)
    _add_header_footer(doc, config)
    _add_title(doc, f"Plano de Aula {int(plan.get('numero', 1)):02d}")

    meta = doc.add_table(rows=6, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(meta)
    _write_labeled_cell(meta.cell(0, 0).merge(meta.cell(0, 3)), "1. CURSO: ", form.get("curso", ""))
    _write_labeled_cell(meta.cell(1, 0).merge(meta.cell(1, 3)), "2. COMPONENTE CURRICULAR: ", form.get("componente", ""))
    _write_labeled_cell(meta.cell(2, 0), "3. MODALIDADE: ", form.get("modalidade", "Presencial"))
    _write_labeled_cell(meta.cell(2, 1), "4. SEMESTRE: ", form.get("semestre_letivo", ""))
    _write_labeled_cell(meta.cell(2, 2), "4.1 PERÍODO: ", form.get("periodo", ""))
    _write_labeled_cell(meta.cell(2, 3), "4.2 TURMA(S): ", form.get("turmas", ""))
    _write_labeled_cell(meta.cell(3, 0).merge(meta.cell(3, 2)), "5. COORDENADOR: ", form.get("coordenador", ""))
    _write_labeled_cell(meta.cell(3, 3), "DURAÇÃO: ", plan.get("duracao", ""))
    _write_labeled_cell(meta.cell(4, 0).merge(meta.cell(4, 2)), "6. TEMA: ", plan.get("tema", ""), fill=PALE_BLUE)
    _write_labeled_cell(meta.cell(4, 3), "6.1 ATIVIDADE: ", plan.get("atividade", ""), fill=PALE_BLUE)
    _write_labeled_cell(meta.cell(5, 0).merge(meta.cell(5, 3)), "DATA: ", plan.get("data", ""))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    objectives = doc.add_table(rows=3, cols=3)
    objectives.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(objectives)
    heading = objectives.cell(0, 0).merge(objectives.cell(0, 2))
    _write_cell(heading, "7. OBJETIVOS / COMPETÊNCIAS", size=8, bold=True, color=WHITE, fill=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for column, (title, key) in enumerate(
        (("7.1 CONHECIMENTOS (Saber)", "conhecimentos"), ("7.2 HABILIDADES (Fazer)", "habilidades"), ("7.3 ATITUDES (Ser)", "atitudes"))
    ):
        _write_cell(objectives.cell(1, column), title, size=8.5, bold=True, color=NAVY, fill=LIGHT_BLUE)
        cell = objectives.cell(2, column)
        _set_cell_margins(cell)
        _clear_paragraph(cell.paragraphs[0])
        _add_bullets(cell, plan.get(key, []), size=8.7, compact=False)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    methods = list(plan.get("metodologia", []))
    resources = list(plan.get("recursos_didaticos", []))
    _add_section_box(doc, "8. Metodologia / recursos didáticos", methods + [f"Recursos: {', '.join(resources)}"], bullets=True)
    _add_section_box(doc, "9. Procedimento avaliativo", plan.get("procedimento_avaliativo", []), bullets=True)
    _add_section_box(doc, "10. Bibliografia", plan.get("bibliografia", []), bullets=True)
    _add_review_notes(doc, plan.get("notas_para_revisao", []))
    return _save(doc)


def build_planos_aula(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    plans = data.get("planos", [])
    if not plans:
        plans = [{
            "numero": 1, "tema": "Plano não gerado", "atividade": "OUTRA", "duracao": "",
            "conhecimentos": [], "habilidades": [], "atitudes": [], "metodologia": [],
            "recursos_didaticos": [], "procedimento_avaliativo": [], "bibliografia": [],
            "notas_para_revisao": data.get("notas_gerais", ["A IA não identificou aulas de conteúdo no cronograma."]),
        }]
    if len(plans) == 1:
        number = int(plans[0].get("numero", 1))
        filename = f"Plano_de_Aula_{number:02d}_{_slug(plans[0].get('tema', 'aula'))}.docx"
        return GeneratedArtifact(
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _build_single_lesson_plan(plans[0], form, config),
        )

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for plan in plans:
            number = int(plan.get("numero", 1))
            filename = f"Plano_de_Aula_{number:02d}_{_slug(plan.get('tema', 'aula'))}.docx"
            archive.writestr(filename, _build_single_lesson_plan(plan, form, config).getvalue())
        if data.get("notas_gerais"):
            archive.writestr("LEIA-ME.txt", "\n".join(data["notas_gerais"]))
    zip_buffer.seek(0)
    filename = f"Planos_de_Aula_{_slug(form.get('componente', 'componente'))}.zip"
    return GeneratedArtifact(filename, "application/zip", zip_buffer)


def build_apostila(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    doc = Document()
    _configure_document(doc, body_size=11)
    _add_header_footer(doc, config)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Cm(3)
    run = cover.add_run("APOSTILA DE AULA")
    _set_run_font(run, size=13, bold=True, color=YELLOW)
    run = cover.add_run(f"\n{data.get('titulo', form.get('titulo_aula', 'Aula'))}")
    _set_run_font(run, size=23, bold=True, color=NAVY)
    run = cover.add_run(f"\n\n{form.get('curso', '')} • {form.get('componente', '')}")
    _set_run_font(run, size=11, color=DARK_GRAY)
    run = cover.add_run(f"\n{form.get('publico_alvo', '')}")
    _set_run_font(run, size=10, italic=True, color=DARK_GRAY)

    note = doc.add_table(rows=1, cols=1)
    note.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(note, color=BLUE, size=6)
    _write_cell(
        note.cell(0, 0),
        "Material de apoio gerado por IA. O professor deve revisar conceitos, referências, exemplos e adequação pedagógica antes da disponibilização.",
        size=9, color=NAVY, fill=LIGHT_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_review_notes(doc, data.get("avisos_validacao", []), "VALIDAÇÃO ANTES DA PUBLICAÇÃO")
    doc.add_page_break()

    p = doc.add_paragraph(data.get("apresentacao", ""))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _format_paragraph(p, size=11, space_after=8, line_spacing=1.15)

    heading = doc.add_paragraph("Objetivos de aprendizagem", style="Heading 1")
    _add_bullets(doc, data.get("objetivos_aprendizagem", []), size=10.5, compact=False)

    for section in data.get("secoes", []):
        doc.add_paragraph(section.get("titulo", "Seção"), style="Heading 1")
        for text in section.get("paragrafos", []):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _format_paragraph(p, size=11, space_after=6, line_spacing=1.15)
        if section.get("pontos_chave"):
            sub = doc.add_paragraph("Pontos-chave", style="Heading 3")
            _add_bullets(doc, section["pontos_chave"], size=10.5, compact=False)
        if section.get("exemplo_aplicado"):
            callout = doc.add_table(rows=1, cols=1)
            _set_table_borders(callout, color=BLUE, size=5)
            cell = callout.cell(0, 0)
            _set_cell_shading(cell, PALE_BLUE)
            _set_cell_margins(cell, 150, 180, 150, 180)
            p = cell.paragraphs[0]
            _clear_paragraph(p)
            run = p.add_run("EXEMPLO APLICADO\n")
            _set_run_font(run, size=9, bold=True, color=NAVY)
            run = p.add_run(section["exemplo_aplicado"])
            _set_run_font(run, size=10.5)
            p.paragraph_format.line_spacing = 1.1

    doc.add_paragraph("Síntese da aula", style="Heading 1")
    _add_bullets(doc, data.get("sintese", []), size=10.5, compact=False)

    doc.add_paragraph("Questões para revisão", style="Heading 1")
    _add_numbered(doc, data.get("questoes_revisao", []), size=10.5)

    if data.get("orientacoes_respostas"):
        doc.add_paragraph("Orientações de resposta", style="Heading 2")
        _add_bullets(doc, data["orientacoes_respostas"], size=10)

    doc.add_paragraph("Referências", style="Heading 1")
    for reference in data.get("referencias", []):
        p = doc.add_paragraph(reference)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        _format_paragraph(p, size=9.5, space_after=5)

    filename = f"Apostila_{_slug(data.get('titulo', form.get('titulo_aula', 'aula')))}.docx"
    return GeneratedArtifact(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _save(doc))


def build_avaliacao(data: dict[str, Any], form: dict[str, Any], config: dict[str, Any]) -> GeneratedArtifact:
    doc = Document()
    _configure_document(doc, body_size=10.5)
    _add_header_footer(doc, config)
    _add_title(doc, data.get("titulo", form.get("titulo_avaliacao", "Avaliação da Aprendizagem")))

    meta = doc.add_table(rows=4, cols=4)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(meta)
    _write_labeled_cell(meta.cell(0, 0).merge(meta.cell(0, 1)), "Curso: ", form.get("curso", ""))
    _write_labeled_cell(meta.cell(0, 2).merge(meta.cell(0, 3)), "Componente: ", form.get("componente", ""))
    _write_labeled_cell(meta.cell(1, 0).merge(meta.cell(1, 2)), "Estudante: ", "")
    _write_labeled_cell(meta.cell(1, 3), "Matrícula: ", "")
    _write_labeled_cell(meta.cell(2, 0), "Turma: ", form.get("turma", ""))
    _write_labeled_cell(meta.cell(2, 1), "Data: ", form.get("data_avaliacao", ""))
    _write_labeled_cell(meta.cell(2, 2), "Valor: ", form.get("valor_total", ""))
    _write_labeled_cell(meta.cell(2, 3), "Tempo: ", form.get("duracao_avaliacao", ""))
    _write_labeled_cell(meta.cell(3, 0).merge(meta.cell(3, 3)), "Professor(a): ", form.get("professor", ""))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    _add_section_box(doc, "Instruções", data.get("instrucoes", []), bullets=True)

    questions = data.get("questoes", [])
    for index, question in enumerate(questions, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(
            f"QUESTÃO {question.get('numero', index):02d} — {question.get('valor', '')} "
            f"| {question.get('nivel_bloom', '')} | {question.get('dificuldade', '')}"
        )
        _set_run_font(run, size=9.5, bold=True, color=NAVY)
        p = doc.add_paragraph(question.get("enunciado", ""))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _format_paragraph(p, size=10.5, space_after=5, line_spacing=1.1)
        alternatives = question.get("alternativas", [])
        for alternative in alternatives:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{alternative.get('letra', '')}) ")
            _set_run_font(run, size=10, bold=True, color=NAVY)
            run = p.add_run(alternative.get("texto", ""))
            _set_run_font(run, size=10)
        if not alternatives:
            for _ in range(6):
                p = doc.add_paragraph("_" * 100)
                _format_paragraph(p, size=8, color="9A9A9A", space_after=2)

    doc.add_page_break()
    _add_title(doc, "Gabarito e Rubricas", "USO EXCLUSIVO DO DOCENTE")
    for index, question in enumerate(questions, start=1):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_borders(table, color="7F8C97", size=5)
        cell = table.cell(0, 0)
        _set_cell_margins(cell, 120, 150, 120, 150)
        _set_cell_shading(cell, PALE_BLUE if index % 2 == 0 else WHITE)
        p = cell.paragraphs[0]
        _clear_paragraph(p)
        run = p.add_run(f"QUESTÃO {question.get('numero', index):02d} — RESPOSTA: ")
        _set_run_font(run, size=9.5, bold=True, color=NAVY)
        run = p.add_run(question.get("resposta_correta", ""))
        _set_run_font(run, size=9.5, bold=True)
        p = cell.add_paragraph()
        run = p.add_run("Justificativa: ")
        _set_run_font(run, size=9, bold=True, color=NAVY)
        run = p.add_run(question.get("justificativa", ""))
        _set_run_font(run, size=9)
        if question.get("rubrica"):
            p = cell.add_paragraph()
            run = p.add_run("Rubrica:")
            _set_run_font(run, size=9, bold=True, color=NAVY)
            _add_bullets(cell, question["rubrica"], size=8.8)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    _add_review_notes(doc, data.get("observacoes_revisao_docente", []), "REVISÃO OBRIGATÓRIA ANTES DA APLICAÇÃO")
    filename = f"Avaliacao_{_slug(form.get('componente', 'componente'))}_{_slug(form.get('turma', 'turma'))}.docx"
    return GeneratedArtifact(filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", _save(doc))


BUILDERS = {
    "ementa": build_ementa,
    "plano-ensino": build_plano_ensino,
    "cronograma": build_cronograma,
    "planos-aula": build_planos_aula,
    "apostila": build_apostila,
    "avaliacao": build_avaliacao,
}


def build_artifact(
    document_type: str,
    generated_data: dict[str, Any],
    form_data: dict[str, Any],
    config: dict[str, Any],
) -> GeneratedArtifact:
    builder = BUILDERS.get(document_type)
    if not builder:
        raise ValueError(f"Construtor não disponível: {document_type}")
    return builder(generated_data, form_data, config)
