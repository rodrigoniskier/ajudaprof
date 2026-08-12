import io

import pytest
from docx import Document
from werkzeug.datastructures import FileStorage

from app.errors import AppError
from app.services import file_processor
from app.services.file_processor import process_upload


def test_extracts_docx_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("Texto principal")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Campo"
    table.cell(0, 1).text = "Valor"
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    storage = FileStorage(stream=buffer, filename="modelo.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    source = process_upload(storage, "Modelo", max_file_bytes=2 * 1024 * 1024)
    assert "Texto principal" in source.text
    assert "Campo | Valor" in source.text
    assert source.inline_bytes is None


def test_pdf_is_kept_inline_in_memory():
    data = b"%PDF-1.4\n% test"
    storage = FileStorage(stream=io.BytesIO(data), filename="arquivo.pdf", content_type="application/pdf")
    source = process_upload(storage, "PDF", max_file_bytes=1024)
    assert source.inline_bytes == data
    assert source.mime_type == "application/pdf"


def test_textual_pdf_can_be_reduced_to_text(monkeypatch):
    monkeypatch.setattr(file_processor, "_extract_pdf", lambda _data: "Texto selecionável do plano de aula")
    data = b"%PDF-1.4\n% test"
    storage = FileStorage(stream=io.BytesIO(data), filename="plano.pdf", content_type="application/pdf")
    source = process_upload(storage, "Plano", max_file_bytes=1024, prefer_pdf_text=True)
    assert source.text == "Texto selecionável do plano de aula"
    assert source.inline_bytes is None


def test_rejects_file_with_fake_pdf_extension():
    storage = FileStorage(stream=io.BytesIO(b"not a pdf"), filename="arquivo.pdf", content_type="application/pdf")
    with pytest.raises(AppError, match="PDF válido"):
        process_upload(storage, "PDF", max_file_bytes=1024)


def test_rejects_disallowed_extension():
    storage = FileStorage(stream=io.BytesIO(b"executable"), filename="arquivo.exe")
    with pytest.raises(AppError, match="Formato não permitido"):
        process_upload(storage, "Anexo", max_file_bytes=1024)
