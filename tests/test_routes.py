import io
import zipfile

from docx import Document

from .conftest import upload


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def assert_docx(response, expected_filename_fragment):
    assert response.status_code == 200, response.get_data(as_text=True)
    assert response.mimetype == DOCX_MIME
    assert response.data[:2] == b"PK"
    assert expected_filename_fragment in response.headers["Content-Disposition"]
    assert float(response.headers["X-Generation-Seconds"]) >= 0
    document = Document(io.BytesIO(response.data))
    assert document.paragraphs or document.tables
    return document


def common(csrf):
    return {"csrf_token": csrf, "privacy_consent": "yes"}


def test_public_pages(client):
    home = client.get("/")
    assert home.status_code == 200
    assert b"Ajuda-Professores" in home.data
    assert b"Prof. Rodrigo Niskier [Medicina - UNIP" in home.data
    assert 'name="av_numero_questoes" type="number" min="1" max="30"' in home.get_data(as_text=True)
    assert "maiores de 18 anos" in home.get_data(as_text=True)
    assert b"Content-Security-Policy" not in home.data
    assert "default-src 'self'" in home.headers["Content-Security-Policy"]
    assert home.headers["Strict-Transport-Security"] == "max-age=31536000; includeSubDomains"
    assert client.get("/privacidade").status_code == 200
    health = client.get("/health").json
    assert health["service"] == "ajuda-professores"
    assert health["status"] == "ok"
    assert health["ai_generation_enabled"] is True
    assert health["fake_mode"] is True
    assert health["ai_configured"] is True


def test_error_page_keeps_institutional_logo(client):
    response = client.get("/rota-inexistente")
    assert response.status_code == 404
    assert b"logo.png" in response.data


def test_ementa_generation(client, csrf):
    data = common(csrf) | {
        "ementa_curso": "Medicina",
        "ementa_componente": "Imunologia Clínica",
        "ementa_carga_horaria": "60 h/a",
        "ementa_periodo": "4º",
        "ementa_observacoes": "Alinhar às DCNs.",
        "ementa_ppc": upload("ppc.txt"),
        "ementa_dcns": upload("dcns.txt"),
    }
    document = assert_docx(client.post("/gerar/ementa", data=data), "Ementa_")
    assert any("EMENTÁRIO" in p.text for p in document.paragraphs)


def test_plano_ensino_generation(client, csrf):
    data = common(csrf) | {
        "pe_curso": "Medicina", "pe_componente": "MAPD I", "pe_periodo": "3º",
        "pe_turno": "Vespertino", "pe_modalidade": "Presencial", "pe_semestre": "2026.2",
        "pe_professor": "Professor responsável", "pe_ch_semanal": "3 h/a", "pe_ch_semestral": "60 h/a",
        "pe_topicos[]": ["Lesão celular", "Imunidade inata", "Microbiologia"],
        "pe_ementa": upload("ementa.txt"),
    }
    document = assert_docx(client.post("/gerar/plano-ensino", data=data), "Plano_de_Ensino_")
    all_text = " ".join(p.text for p in document.paragraphs) + " ".join(cell.text for t in document.tables for row in t.rows for cell in row.cells)
    assert "PLANO DE ENSINO" in all_text


def test_cronograma_generation(client, csrf):
    data = common(csrf) | {
        "cr_curso": "Medicina", "cr_componente": "MAPD II", "cr_docentes": "Docente A\nDocente B",
        "cr_coordenador": "Coordenação", "cr_periodo": "4º", "cr_semestre": "2026.2", "cr_turma": "A",
        "cr_ch_teorica": "30 h/a", "cr_ch_pratica": "30 h/a", "cr_data_inicio": "2026-08-03",
        "cr_data_fim": "2026-12-15", "cr_dia[]": ["Segunda-feira", "Quarta-feira"],
        "cr_inicio[]": ["13:00", "13:00"], "cr_fim[]": ["15:30", "15:30"],
        "cr_docente_dia[]": ["Docente A", "Docente B"],
        "cr_topicos[]": ["Tópico 1", "Tópico 2", "Tópico 3"],
        "cr_calendario": upload("calendario.txt"),
    }
    document = assert_docx(client.post("/gerar/cronograma", data=data), "Cronograma_")
    assert document.sections[0].orientation is not None


def test_single_lesson_plan_generation(client, csrf):
    data = common(csrf) | {
        "pa_curso": "Medicina", "pa_componente": "MAPD II", "pa_modalidade": "Presencial",
        "pa_semestre": "2026.2", "pa_periodo": "4º", "pa_turmas": "A, B", "pa_coordenador": "Coordenação",
        "pa_aula_alvo": "Aula 1", "pa_cronograma": upload("cronograma.txt"),
        "pa_dcns": upload("dcns.txt"), "pa_ementa": upload("ementa.txt"),
    }
    document = assert_docx(client.post("/gerar/planos-aula", data=data), "Plano_de_Aula_")
    assert any("PLANO DE AULA" in p.text for p in document.paragraphs)


def test_apostila_generation(client, csrf):
    data = common(csrf) | {
        "ap_curso": "Medicina", "ap_componente": "MAPD II", "ap_titulo": "Imunidade inata",
        "ap_publico": "4º período", "ap_nivel": "Intermediário", "ap_extensao": "Regular — 7 a 12 páginas",
        "ap_plano": upload("plano.txt"), "ap_referencias": upload("referencia.txt"),
    }
    document = assert_docx(client.post("/gerar/apostila", data=data), "Apostila_")
    assert any("APOSTILA DE AULA" in p.text for p in document.paragraphs)


def test_evaluation_generation(client, csrf):
    data = common(csrf) | {
        "av_curso": "Medicina", "av_componente": "MAPD II", "av_titulo": "Avaliação 1",
        "av_professor": "Docente", "av_turma": "A", "av_data": "2026-09-01", "av_duracao": "90 min",
        "av_valor_total": "10,0 pontos", "av_numero_questoes": "2",
        "av_tipo[]": ["Múltipla escolha — resposta única", "Discursiva"],
        "av_quantidade[]": ["1", "1"], "av_valor[]": ["1,0", "2,0"],
        "av_alternativas[]": ["4", ""], "av_outro[]": ["", ""],
        "av_bloom[]": ["Compreender", "Aplicar", "Analisar"],
        "av_facil": "20", "av_media": "60", "av_dificil": "20",
        "av_contextualizacao": "Predominantemente contextualizada", "av_somente_anexos": "yes",
        "av_apostilas": upload("apostila.txt"), "av_planos": upload("plano.txt"),
    }
    document = assert_docx(client.post("/gerar/avaliacao", data=data), "Avaliacao_")
    assert any("GABARITO" in p.text.upper() for p in document.paragraphs)


def test_evaluation_rejects_more_than_thirty_questions(client, csrf):
    response = client.post(
        "/gerar/avaliacao",
        data=common(csrf) | {"av_numero_questoes": "31"},
        headers={"Accept": "application/json"},
    )
    assert response.status_code == 400
    assert "entre 1 e 30" in response.json["message"]


def test_privacy_consent_required(client, csrf):
    response = client.post("/gerar/ementa", data={"csrf_token": csrf}, headers={"Accept": "application/json"})
    assert response.status_code == 400
    assert response.json["error"] == "validation_error"


def test_invalid_csrf_rejected(client):
    response = client.post(
        "/gerar/ementa",
        data={"csrf_token": "invalid", "privacy_consent": "yes"},
        headers={"Accept": "application/json"},
    )
    assert response.status_code == 400
    assert response.json["error"] == "csrf_invalid"


def test_generation_disabled_by_kill_switch(app):
    app.config["AI_GENERATION_ENABLED"] = False
    client = app.test_client()
    response = client.post("/gerar/ementa", headers={"Accept": "application/json"})
    assert response.status_code == 503
    assert response.json["error"] == "generation_disabled"
    assert client.get("/health").json["ai_generation_enabled"] is False
    # Páginas estáticas continuam no ar (degradação graciosa).
    assert client.get("/").status_code == 200
